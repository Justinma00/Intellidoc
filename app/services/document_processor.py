import os
import uuid
from typing import Dict, Any
from pathlib import Path
import PyPDF2
import docx
from ..config import settings

class DocumentProcessor:
    def __init__(self) -> None:
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
    def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file and return file path."""
        # Generate unique filename
        file_ext = Path(filename).suffix
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = self.upload_dir / unique_filename
        
        with open(file_path, "wb") as f:
            f.write(file_content)
            
        return str(file_path)
    
    def extract_text_from_file(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Extract text from various file types."""
        try:
            if mime_type == "application/pdf":
                return self._extract_from_pdf(file_path)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_from_docx(file_path)
            elif mime_type.startswith("image/"):
                return self._extract_from_image(file_path)
            elif mime_type == "text/plain":
                return self._extract_from_text(file_path)
            else:
                return {
                    "text": "",
                    "error": f"Unsupported file type: {mime_type}"
                }
        except Exception as e:
            return {
                "text": "",
                "error": f"Error extracting text: {str(e)}"
            }
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            return {
                "text": text.strip(),
                "pages": len(pdf_reader.pages),
                "method": "pdf_extraction"
            }
        except Exception as e:
            return {
                "text": "",
                "error": f"PDF extraction error: {str(e)}"
            }
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX."""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "text": text.strip(),
                "paragraphs": len(doc.paragraphs),
                "method": "docx_extraction"
            }
        except Exception as e:
            return {
                "text": "",
                "error": f"DOCX extraction error: {str(e)}"
            }
    
    def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from image using OCR."""
        try:
            import cv2  # type: ignore
            import numpy as np  # type: ignore
            import pytesseract  # type: ignore
            # Load and preprocess image
            image = cv2.imread(file_path)
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Apply image processing for better OCR
            gray = cv2.medianBlur(gray, 3)
            gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(gray)
            
            return {
                "text": text.strip(),
                "method": "ocr_extraction",
                "image_size": image.shape[:2]
            }
        except Exception as e:
            return {
                "text": "",
                "error": f"OCR extraction error: {str(e)}"
            }
    
    def _extract_from_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                "text": text.strip(),
                "method": "text_file"
            }
        except Exception as e:
            return {
                "text": "",
                "error": f"Text extraction error: {str(e)}"
            }
    
    def analyze_document_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure and extract metadata."""
        lines = text.split('\n')
        words = text.split()
        
        # Basic statistics
        stats = {
            "total_characters": len(text),
            "total_words": len(words),
            "total_lines": len(lines),
            "average_words_per_line": len(words) / max(len(lines), 1),
            "average_characters_per_word": len(text) / max(len(words), 1)
        }
        
        # Extract potential headers (lines with fewer words)
        potential_headers = [
            line.strip() for line in lines 
            if line.strip() and len(line.split()) <= 5 and line.strip().isupper()
        ]
        
        # Extract potential dates
        import re
        date_pattern = r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'
        dates = re.findall(date_pattern, text)
        
        # Extract potential email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        
        # Extract potential phone numbers
        phone_pattern = r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
        phones = re.findall(phone_pattern, text)
        
        return {
            "statistics": stats,
            "potential_headers": potential_headers[:10],  # Limit to first 10
            "dates_found": dates[:5],  # Limit to first 5
            "emails_found": emails[:5],  # Limit to first 5
            "phones_found": phones[:5],  # Limit to first 5
        }